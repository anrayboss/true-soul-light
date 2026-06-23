import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_runs_with_bold(paragraph, text):
    # Split text by '**' to find bold segments
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set standard margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_list = False
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            # Add space or empty paragraph if needed, but in docx empty lines between items are usually handled by paragraph spacing.
            continue
            
        # Check for headings
        if line_stripped.startswith('# '):
            text = line_stripped[2:]
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            in_list = False
        elif line_stripped.startswith('## '):
            text = line_stripped[3:]
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            in_list = False
        elif line_stripped.startswith('### '):
            text = line_stripped[4:]
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            in_list = False
        # Check for list items
        elif line.startswith('- ') or line.startswith('* '):
            # Level 1 list
            text = line.lstrip('-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_bold(p, text)
            p.paragraph_format.space_after = Pt(3)
            in_list = True
        elif line.startswith('  - ') or line.startswith('  * '):
            # Level 2 list
            text = line.strip().lstrip('-* ').strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_runs_with_bold(p, text)
            p.paragraph_format.space_after = Pt(3)
            in_list = True
        elif line.startswith('    - ') or line.startswith('    * '):
            # Level 3 list
            text = line.strip().lstrip('-* ').strip()
            p = doc.add_paragraph(style='List Bullet 3')
            add_runs_with_bold(p, text)
            p.paragraph_format.space_after = Pt(3)
            in_list = True
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            add_runs_with_bold(p, line_stripped)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            in_list = False
            
    # Apply standard styling to all fonts (Calibri for english, Microsoft JhengHei for Chinese)
    # We can set the font properties on the 'Normal' style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    
    # We also want to set Heading fonts
    for lvl in range(1, 4):
        h_style = doc.styles[f'Heading {lvl}']
        h_font = h_style.font
        h_font.name = 'Microsoft JhengHei'
        h_font.color.rgb = None  # Use default or automatic color
        
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python convert_md_to_docx.py <input_md> <output_docx>")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
