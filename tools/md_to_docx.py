import re
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def parse_inline(paragraph, text):
    """
    Parse inline formats like **bold**, *italic*, `code` in text and add to paragraph.
    """
    # Regex to tokenise bold, italic, code and normal text
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)

def set_cell_background(cell, color_hex):
    """
    Set background color of a table cell.
    """
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_docx(md_path, docx_path):
    doc = Document()
    
    # Configure styles (normal font: Microsoft JhengHei for Traditional Chinese)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_rows = []
    
    in_code_block = False
    code_block_text = ""
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(code_block_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(80, 80, 80)
                p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="4" w:color="CCCCCC"/></w:pBdr>')
                p._p.get_or_add_pPr().append(p_border)
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._p.get_or_add_pPr().append(shading_elm)
                
                in_code_block = False
                code_block_text = ""
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_block_text += line
            i += 1
            continue
            
        # Table handling
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # We finished reading a table block, process it
            process_table(doc, table_rows)
            table_rows = []
            in_table = False
            
        if not stripped:
            i += 1
            continue
            
        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            run.font.name = 'Microsoft JhengHei'
            run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
            run.bold = True
        elif stripped.startswith('## '):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            run.font.name = 'Microsoft JhengHei'
            run.font.color.rgb = RGBColor(0, 102, 153)
            run.bold = True
        elif stripped.startswith('### '):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[4:])
            run.font.name = 'Microsoft JhengHei'
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.bold = True
        # Bullet list
        elif stripped.startswith('* ') or stripped.startswith('- '):
            # Determine indentation level
            indent_level = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent_level // 2))
            parse_inline(p, stripped[2:])
        # Numbered list
        elif re.match(r'^\d+\.\s+', stripped):
            match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            num = match.group(1)
            content = match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, content)
        # Horizontal rule
        elif stripped in ['---', '***']:
            p = doc.add_paragraph()
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="DDDDDD"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parse_inline(p, stripped)
            
        i += 1
        
    if in_table:
        process_table(doc, table_rows)
        
    doc.save(docx_path)
    print(f"Document saved successfully: {docx_path}")

def process_table(doc, rows):
    """
    Process markdown table rows and insert a formatted Word table.
    """
    if not rows:
        return
        
    parsed_rows = []
    for r in rows:
        cells = [c.strip() for c in r.strip().split('|')]
        if r.strip().startswith('|'):
            cells = cells[1:]
        if r.strip().endswith('|') and len(cells) > 0:
            cells = cells[:-1]
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return
        
    headers = parsed_rows[0]
    has_separator = False
    if len(parsed_rows) > 1:
        sec_row = parsed_rows[1]
        if all(re.match(r'^:?-+:?$', c) for c in sec_row):
            has_separator = True
            
    data_rows = parsed_rows[2:] if has_separator else parsed_rows[1:]
    
    num_cols = len(headers)
    num_rows = len(data_rows) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = text
        set_cell_background(hdr_cells[col_idx], "4F81BD")
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.name = 'Microsoft JhengHei'
            
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F2F5F8" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx in range(num_cols):
            if col_idx < len(row_data):
                text = row_data[col_idx]
                row_cells[col_idx].text = text
                p = row_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for r in p.runs:
                    r.font.name = 'Microsoft JhengHei'
                    r.font.size = Pt(10)
            set_cell_background(row_cells[col_idx], bg_color)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input_markdown_file> [output_docx_file]")
        sys.exit(1)
        
    input_md = sys.argv[1]
    if len(sys.argv) >= 3:
        output_docx = sys.argv[2]
    else:
        output_docx = os.path.splitext(input_md)[0] + ".docx"
        
    try:
        import pypandoc
        print("Attempting conversion using pypandoc...")
        pypandoc.convert_file(input_md, 'docx', outputfile=output_docx)
        print(f"Document saved successfully via pypandoc: {output_docx}")
    except Exception as e:
        print(f"pypandoc conversion failed or not available ({e}). Falling back to python-docx...")
        create_docx(input_md, output_docx)

